"""Shared fixtures for tests."""

from pathlib import Path

import pytest
from docx import Document


GOOD_RESUME = """\
John Doe
john.doe@email.com | (555) 123-4567 | linkedin.com/in/johndoe

Professional Summary
Results-driven software engineer with 5+ years of experience building scalable web apps.

Experience
Senior Software Engineer - Acme Corp (2020-2025)
- Led a team of 8 engineers to deliver a microservices platform, reducing deploy time by 40%.
- Designed and implemented RESTful APIs serving 10M+ requests/day.
- Improved CI/CD pipeline efficiency by 35%, saving $50K annually.
- Built real-time monitoring dashboards used by 200+ developers.
- Managed migration of legacy monolith to cloud-native architecture on AWS.
- Optimized database queries, cutting p99 latency from 800ms to 120ms.
- Mentored 4 junior engineers through structured onboarding programme.

Education
B.S. Computer Science - MIT (2016-2020)

Skills
Python, JavaScript, TypeScript, React, Node.js, AWS, Docker, Kubernetes, PostgreSQL, Redis

Projects
Open-source CLI tool for automated code review (500+ GitHub stars).

Certifications
AWS Solutions Architect - Associate
"""

BARE_RESUME = "Hello world, this is a very bare resume with no structure."


@pytest.fixture()
def good_resume_text() -> str:
    return GOOD_RESUME


@pytest.fixture()
def bare_resume_text() -> str:
    return BARE_RESUME


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    """Create a real .docx file on disk."""
    filepath = tmp_path / "resume.docx"
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("jane@example.com | (555) 999-0000")
    doc.add_paragraph("Experience")
    doc.add_paragraph("- Built an API that handled 1M requests.")
    doc.save(str(filepath))
    return filepath
